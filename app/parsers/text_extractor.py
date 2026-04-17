import os
from io import BytesIO
from pathlib import Path
import subprocess
import tempfile


def extract_text(filename: str, payload: bytes) -> str:
    suffix = filename.rsplit(".", maxsplit=1)[-1].lower() if "." in filename else ""

    if suffix in {"txt", "log"}:
        return payload.decode("utf-8")

    if suffix == "pdf":
        return _extract_pdf_text(payload)

    if suffix == "docx":
        return _extract_docx_text(payload)

    if suffix == "doc":
        return _extract_doc_text(payload)

    raise ValueError("Unsupported file type")


def _extract_pdf_text(payload: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF processing dependency is not installed") from exc

    reader = PdfReader(BytesIO(payload))
    content = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    if not content:
        raise ValueError("PDF does not contain extractable text")
    return content


def _extract_docx_text(payload: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Word processing dependency is not installed") from exc

    document = Document(BytesIO(payload))
    segments = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    segments.append(cell_text)

    content = "\n".join(segments).strip()
    if not content:
        raise ValueError("Word document does not contain extractable text")
    return content


def _extract_doc_text(payload: bytes) -> str:
    if os.name != "nt":
        raise ValueError("Legacy .doc files are only supported on Windows")

    with tempfile.TemporaryDirectory(prefix="ewa-doc-") as temp_dir:
        source_path = Path(temp_dir) / "input.doc"
        target_path = Path(temp_dir) / "converted.docx"
        source_path.write_bytes(payload)

        _convert_doc_to_docx(source_path, target_path)
        return _extract_docx_text(target_path.read_bytes())


def _convert_doc_to_docx(source_path: Path, target_path: Path) -> None:
    # 16 is the Word constant for wdFormatDocumentDefault (.docx).
    script = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $document = $word.Documents.Open('{_escape_powershell_path(source_path)}')
    $document.SaveAs([ref] '{_escape_powershell_path(target_path)}', [ref] 16)
}} finally {{
    if ($document -ne $null) {{ $document.Close() }}
    if ($word -ne $null) {{ $word.Quit() }}
}}
""".strip()

    result = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0 or not target_path.exists():
        raise ValueError("Microsoft Word is required to process legacy .doc files")


def _escape_powershell_path(path: Path) -> str:
    return str(path).replace("'", "''")
