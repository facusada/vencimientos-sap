import os
import sys
from io import BytesIO
from pathlib import Path, PosixPath, WindowsPath
import shutil
import subprocess
import tempfile
import xml.etree.ElementTree as ET

from app.services.ocr_service import create_ocr_service
from app.services.ocr_service import render_pdf_pages_to_images

def extract_text(filename: str, payload: bytes) -> str:
    suffix = filename.rsplit(".", maxsplit=1)[-1].lower() if "." in filename else ""

    if suffix in {"txt", "log"}:
        return payload.decode("utf-8")

    if suffix == "pdf":
        return _extract_pdf_text(payload)

    if suffix == "docx":
        return _extract_docx_text(payload)

    if suffix == "doc":
        return _extract_doc_text(payload)

    raise ValueError("Unsupported file type")


def _extract_pdf_text(payload: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF processing dependency is not installed") from exc

    reader = PdfReader(BytesIO(payload))
    content = "\n".join(page.extract_text() or "" for page in reader.pages).strip()
    ocr_content = _extract_pdf_ocr_text(payload)
    if ocr_content:
        content = "\n".join(part for part in [content, ocr_content] if part).strip()
    if not content:
        raise ValueError("PDF does not contain extractable text")
    return content


def _extract_docx_text(payload: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Word processing dependency is not installed") from exc

    document = Document(BytesIO(payload))
    segments = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    segments.append(cell_text)

    ocr_content = _extract_docx_ocr_text(payload)
    if ocr_content:
        segments.append(ocr_content)

    content = "\n".join(segments).strip()
    if not content:
        raise ValueError("Word document does not contain extractable text")
    return content


def _extract_doc_text(payload: bytes) -> str:
    if _looks_like_word_2003_xml(payload):
        content = _extract_word_2003_xml_text(payload)
        ocr_content = _extract_office_document_ocr_text(payload, "doc")
        if ocr_content:
            content = "\n".join(part for part in [content, ocr_content] if part).strip()
        return content

    with tempfile.TemporaryDirectory(prefix="ewa-doc-") as temp_dir:
        source_path = _build_os_path(temp_dir) / "input.doc"
        target_path = _build_os_path(temp_dir) / "input.docx"
        source_path.write_bytes(payload)

        _convert_doc_to_docx(source_path, target_path)
        return _extract_docx_text(target_path.read_bytes())


def _convert_doc_to_docx(source_path: Path, target_path: Path) -> None:
    if os.name == "nt":
        _convert_doc_to_docx_windows(source_path, target_path)
        return

    _convert_doc_to_docx_non_windows(source_path, target_path)


def _convert_doc_to_docx_windows(source_path: Path, target_path: Path) -> None:
    # 16 is the Word constant for wdFormatDocumentDefault (.docx).
    script = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $document = $word.Documents.Open('{_escape_powershell_path(source_path)}')
    $document.SaveAs([ref] '{_escape_powershell_path(target_path)}', [ref] 16)
}} finally {{
    if ($document -ne $null) {{ $document.Close() }}
    if ($word -ne $null) {{ $word.Quit() }}
}}
""".strip()

    result = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0 or not target_path.exists():
        raise ValueError("Microsoft Word is required to process legacy .doc files")


def _convert_doc_to_docx_non_windows(source_path: Path, target_path: Path) -> None:
    libreoffice_binary = shutil.which("soffice") or shutil.which("libreoffice")
    if libreoffice_binary is None:
        raise ValueError("LibreOffice is required to process legacy .doc files on macOS/Linux")

    result = subprocess.run(
        [
            libreoffice_binary,
            "--headless",
            "--convert-to",
            "docx",
            "--outdir",
            str(target_path.parent),
            str(source_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0 or not target_path.exists():
        raise ValueError("LibreOffice is required to process legacy .doc files on macOS/Linux")


def _looks_like_word_2003_xml(payload: bytes) -> bool:
    leading_text = payload[:2048].decode("utf-8", errors="ignore").lower()
    return (
        "<?xml" in leading_text
        and "progid=\"word.document\"" in leading_text
        and "schemas.microsoft.com/office/word/2003/wordml" in leading_text
    )


def _extract_word_2003_xml_text(payload: bytes) -> str:
    try:
        root = ET.fromstring(payload)
    except ET.ParseError as exc:
        raise ValueError("Word XML document is not valid XML") from exc

    segments: list[str] = []
    seen: set[str] = set()

    for node in root.iter():
        if not node.tag.endswith("}t"):
            continue

        text = "".join(node.itertext()).strip()
        if not text or text in seen:
            continue

        seen.add(text)
        segments.append(text)

    content = "\n".join(segments).strip()
    if not content:
        raise ValueError("Word document does not contain extractable text")
    return content


def _escape_powershell_path(path: Path) -> str:
    return str(path).replace("'", "''")


def _extract_pdf_ocr_text(payload: bytes) -> str:
    try:
        images = render_pdf_pages_to_images(payload)
        return create_ocr_service().extract_text_from_images(images)
    except Exception:
        return ""


def _extract_docx_ocr_text(payload: bytes) -> str:
    return _extract_office_document_ocr_text(payload, "docx")


def _extract_office_document_ocr_text(payload: bytes, suffix: str) -> str:
    with tempfile.TemporaryDirectory(prefix=f"ewa-{suffix}-ocr-") as temp_dir:
        source_path = _build_os_path(temp_dir) / f"input.{suffix}"
        target_path = _build_os_path(temp_dir) / "input.pdf"
        source_path.write_bytes(payload)

        try:
            _convert_office_document_to_pdf(source_path, target_path)
        except ValueError:
            return ""

        if not target_path.exists():
            return ""

        return _extract_pdf_ocr_text(target_path.read_bytes())


def _convert_office_document_to_pdf(source_path: Path, target_path: Path) -> None:
    if os.name == "nt":
        _convert_office_document_to_pdf_windows(source_path, target_path)
        return

    _convert_office_document_to_pdf_non_windows(source_path, target_path)


def _convert_office_document_to_pdf_windows(source_path: Path, target_path: Path) -> None:
    # 17 is the Word constant for wdFormatPDF.
    script = f"""
$ErrorActionPreference = 'Stop'
$word = $null
$document = $null
try {{
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $document = $word.Documents.Open('{_escape_powershell_path(source_path)}')
    $document.SaveAs([ref] '{_escape_powershell_path(target_path)}', [ref] 17)
}} finally {{
    if ($document -ne $null) {{ $document.Close() }}
    if ($word -ne $null) {{ $word.Quit() }}
}}
""".strip()

    result = subprocess.run(
        ["powershell", "-NoProfile", "-NonInteractive", "-Command", script],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0 or not target_path.exists():
        raise ValueError("Microsoft Word is required to render Word documents for OCR")


def _convert_office_document_to_pdf_non_windows(source_path: Path, target_path: Path) -> None:
    libreoffice_binary = shutil.which("soffice") or shutil.which("libreoffice")
    if libreoffice_binary is None:
        raise ValueError("LibreOffice is required to render Word documents for OCR on macOS/Linux")

    result = subprocess.run(
        [
            libreoffice_binary,
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(target_path.parent),
            str(source_path),
        ],
        capture_output=True,
        text=True,
        check=False,
    )

    if result.returncode != 0 or not target_path.exists():
        raise ValueError("LibreOffice is required to render Word documents for OCR on macOS/Linux")


def _build_os_path(base_path: str):
    path_cls = WindowsPath if sys.platform.startswith("win") else PosixPath
    return path_cls(base_path)
