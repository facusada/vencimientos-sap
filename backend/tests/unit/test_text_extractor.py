from io import BytesIO
from pathlib import Path
import shutil
import subprocess

import pytest
from docx import Document

from app.parsers.text_extractor import _extract_doc_text
from app.parsers.text_extractor import extract_text


def test_extract_text_supports_docx_files():
    document = Document()
    document.add_paragraph("Kernel expiry: 2026-12-31")
    document.add_paragraph("Certificate ABC valid until 31.12.2026")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "SAP Product Version"
    table.cell(0, 1).text = "supported until 02.2027"

    buffer = BytesIO()
    document.save(buffer)

    result = extract_text("ewa.docx", buffer.getvalue())

    assert "Kernel expiry: 2026-12-31" in result
    assert "Certificate ABC valid until 31.12.2026" in result
    assert "SAP Product Version" in result
    assert "supported until 02.2027" in result


def test_extract_text_supports_doc_files_via_legacy_extractor(monkeypatch: pytest.MonkeyPatch):
    def fake_extract_doc_text(payload: bytes) -> str:
        assert payload == b"legacy-doc"
        return "Kernel expiry: 2026-12-31"

    monkeypatch.setattr("app.parsers.text_extractor._extract_doc_text", fake_extract_doc_text)

    result = extract_text("ewa.doc", b"legacy-doc")

    assert result == "Kernel expiry: 2026-12-31"


def test_extract_doc_text_supports_word_2003_xml_without_external_converter():
    payload = b"""<?xml version="1.0"?>
<?mso-application progid="Word.Document"?>
<w:wordDocument xmlns:w="http://schemas.microsoft.com/office/word/2003/wordml">
  <w:body>
    <w:p><w:r><w:t>EHP7 FOR SAP ERP 6.0</w:t></w:r></w:p>
    <w:p><w:r><w:t>31.12.2027</w:t></w:r></w:p>
    <w:tbl>
      <w:tr>
        <w:tc><w:p><w:r><w:t>SAP NETWEAVER 7.4</w:t></w:r></w:p></w:tc>
        <w:tc><w:p><w:r><w:t>31.12.2027</w:t></w:r></w:p></w:tc>
      </w:tr>
    </w:tbl>
  </w:body>
</w:wordDocument>
"""

    result = _extract_doc_text(payload)

    assert "EHP7 FOR SAP ERP 6.0" in result
    assert "31.12.2027" in result
    assert "SAP NETWEAVER 7.4" in result


def test_extract_doc_text_supports_non_windows_via_libreoffice(monkeypatch: pytest.MonkeyPatch):
    monkeypatch.setattr("app.parsers.text_extractor.os.name", "posix")
    monkeypatch.setattr(shutil, "which", lambda name: "/usr/bin/soffice")

    def fake_run(command: list[str], capture_output: bool, text: bool, check: bool) -> subprocess.CompletedProcess[str]:
        assert command[:4] == ["/usr/bin/soffice", "--headless", "--convert-to", "docx"]
        outdir = Path(command[command.index("--outdir") + 1])
        converted = outdir / "input.docx"
        converted.write_bytes(b"converted-docx")
        return subprocess.CompletedProcess(command, 0, stdout="", stderr="")

    monkeypatch.setattr(subprocess, "run", fake_run)
    monkeypatch.setattr(
        "app.parsers.text_extractor._extract_docx_text",
        lambda payload: "Kernel expiry: 2026-12-31" if payload == b"converted-docx" else "",
    )

    result = _extract_doc_text(b"legacy-doc")

    assert result == "Kernel expiry: 2026-12-31"


def test_extract_doc_text_requires_supported_converter_on_non_windows(
    monkeypatch: pytest.MonkeyPatch,
):
    monkeypatch.setattr("app.parsers.text_extractor.os.name", "posix")
    monkeypatch.setattr(shutil, "which", lambda name: None)

    with pytest.raises(ValueError, match="LibreOffice is required to process legacy .doc files"):
        _extract_doc_text(b"legacy-doc")
