from io import BytesIO

import pytest
from docx import Document

from app.parsers.text_extractor import extract_text


def test_extract_text_supports_docx_files():
    document = Document()
    document.add_paragraph("Kernel expiry: 2026-12-31")
    document.add_paragraph("Certificate ABC valid until 31.12.2026")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "SAP Product Version"
    table.cell(0, 1).text = "supported until 02.2027"

    buffer = BytesIO()
    document.save(buffer)

    result = extract_text("ewa.docx", buffer.getvalue())

    assert "Kernel expiry: 2026-12-31" in result
    assert "Certificate ABC valid until 31.12.2026" in result
    assert "SAP Product Version" in result
    assert "supported until 02.2027" in result


def test_extract_text_supports_doc_files_via_legacy_extractor(monkeypatch: pytest.MonkeyPatch):
    def fake_extract_doc_text(payload: bytes) -> str:
        assert payload == b"legacy-doc"
        return "Kernel expiry: 2026-12-31"

    monkeypatch.setattr("app.parsers.text_extractor._extract_doc_text", fake_extract_doc_text)

    result = extract_text("ewa.doc", b"legacy-doc")

    assert result == "Kernel expiry: 2026-12-31"
